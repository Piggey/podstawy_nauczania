from docx import Document
from docx.shared import Pt
from json import load
from sys import platform
try:
    from os import startfile
except ImportError as e: # no startfile when on linux
    pass

def getLines():
    # get the input
    ## Sample input:
    # I. 1. 1) 2) 3) 5) 2. 1) 2) 3) 4) 7) 3. 1) 2) 3) 4) 4. 1) 4) 5) 7) 8) 5. 2) 4) 5) 6)
    # II. 4. 2), 6. 8) 9)
    # III. 1. 1) 3) 4)
    lastLine = 'line'
    lines = []
    print('Wklej indeksy rozdzialow tutaj: ')
    while True:
        line = input()
        line = line.replace(', ', ' ').replace(',', ' ') # ignore ',' character and add whitespace in case of input "1),2)"
        
        # whitespace buffer due to MS Word acting weird when copying 
        if(not lastLine and not line):
            break
        else:
            lines.append(line)

        lastLine = line

    return lines

def parseInput(lines):
    # parse input into dictionary calls
    # edu
    ## zakres
    ### osiagniecia
    calls = [] # list of dictionaries

    for line in lines:
        line = line.split(' ')
        edu = line[0]
        for index in line[1:]:
            if(not index.endswith(')')):
                zakres = index
                continue
            else:
                osiagniecie = index

            tmp = {
                'edu': edu,
                'zakres': zakres,
                'osiagniecie': osiagniecie
            }
            # print(tmp)
            calls.append(tmp)
    
    return calls

def createDocument(data, lines):
    ## create .docx document
    # header: wklejony input, Times New Roman, 10
    # edu: Times New Roman, 12, pogrubione
    # zakres: Times New Roman, 12, podkreslenie
    # osiagniecie: Times New Roman, 12, zwykly

    print("tworzenie dokumentu...")
    calls = parseInput(lines)
    d = Document()
    font = d.styles['Normal'].font
    font.name = 'Times New Roman'
    font.size = Pt(10)

    # # create header
    # for line in lines:
    #     if(line): # ignore empty lines
    #         p = d.add_paragraph(line)

    font.size = Pt(12)
    last_edu = ''
    last_zakres = ''
    for call in calls:
        
        # create edu section
        edu = call['edu']
        if(last_edu != edu):
            p = d.add_paragraph()
            try: # key error handlers because i dont like when it just quits on me
                text = '\n' + edu + ' ' + data[edu]['nazwa']
            except KeyError as e:
                endProg(e)

            p.add_run(text).bold = True

        # create zakres section
        zakres = call['zakres']
        if(last_zakres != zakres):
            p = d.add_paragraph() 
            try:
                text = '\n' + zakres + ' ' + data[edu][zakres]['nazwa']
            except KeyError as e:
                endProg(e)

            p.add_run(text).underline = True

        # create osiagniecie section
        osiagniecie = call['osiagniecie']
        try:
            text = osiagniecie + ' ' + data[edu][zakres][osiagniecie]
        except KeyError as e:
            endProg(e)

        # d.add_paragraph(text)
        d.add_paragraph(text)

        last_edu = edu
        last_zakres = zakres

    filename = 'podstawy_programowe.docx'
    d.save(filename)
    try:
        if(platform == 'win32'):
            startfile(filename)
    except PermissionError:
        print("Poprzedni plik jest nadal otwarty! Zamknij go przed uzyciem skryptu.")

def endProg(e):
    print('UWAGA: wykryto bledny rozdzial w podanym tekscie: ' + str(e))
    print('--------')
    input("Nacisnij ENTER, aby zakonczyc...")
    quit(1)

# -- main -- #

# load json file
with open('podstawy_programowe.json', encoding='utf-8') as f:
    data = load(f)

lines = getLines()
createDocument(data, lines)
